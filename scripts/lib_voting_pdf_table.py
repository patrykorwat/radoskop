"""Parser tabel imiennych głosowań z PDFów formatu eSesja.

Reusable biblioteka dla sejmików i rad miast które publikują wyniki głosowań
w PDFach generowanych przez app.esesja.pl. Format "eSesja standard":

- gdański (radoskop/scripts/parse_pdf.py używa tego samego formatu, ta lib
  jest jego refaktor do reusable formy)
- lubelski Sejmik (potwierdzone, 2026-05-19, 16/17 OK)
- świętokrzyski Sejmik (potwierdzone, 2026-05-19, 11/11 OK)

Format PDF eSesja standard:
    [opcjonalny header: nagłówek sesji, data, kadencja]
    Głosowano w sprawie: {tytuł uchwały, może być wieloliniowy}
    Wyniki głosowania
    ZA: N, PRZECIW: N, WSTRZYMUJĘ SIĘ: N, BRAK GŁOSU: N, NIEOBECNI: N
    Wyniki imienne:
    ZA (N) {nazwiska oddzielone przecinkami}
    PRZECIW (N) ...
    WSTRZYMUJĘ SIĘ (N) ...
    BRAK GŁOSU (N) ...
    NIEOBECNI (N) ...
    Głosowanie z dnia: DD.MM.YYYY, HH:MM:SS
    Wygenerowano za pomocą app.esesja.pl

INNE FORMATY (NIE pokrywa ta lib):

* eSesja "wydruk per strona" (kujawsko-pomorskie, prawdopodobnie inne) -
  1 głosowanie = 1 strona, dwukolumnowa tabela "Lp | Nazwisko | Decyzja"
  + header "Typ głosowania" + "Data głosowania". Brak frazy "Głosowano
  w sprawie:" i "Wyniki imienne:". Wymaga osobnego parsera.
* PDF-as-image (skanowane, kujawsko-pomorskie, zachodniopomorskie) -
  wymaga OCR przez `extract_pdf_text(ocr_fallback=True)`. OCR po polsku
  wymaga `tesseract-ocr-pol` (apt). Bez polish pack accuracy spada bo
  diakrytyki giną.
* Protokoły narracyjne (lubuskie) - PDF zawiera prozę protokołu, nie
  tabelę głosowań. Wymaga osobnego źródła danych albo NLP.

Detekcja formatu: użyj `detect_pdf_format(pdf_path)` żeby się dowiedzieć
która ścieżka parsing będzie działać.

Usage:
    from lib_voting_pdf_table import parse_voting_pdf, detect_pdf_format

    fmt = detect_pdf_format("session.pdf")
    if fmt == "esesja_standard":
        result = parse_voting_pdf("session.pdf")
        print(result["vote_count"])
        for v in result["votes"]:
            print(v["topic"][:60], v["counts"])
            print("  ZA:", v["named_votes"]["za"][:3])
    elif fmt == "scanned":
        result = parse_voting_pdf("session.pdf")  # OCR fallback automatic
    elif fmt == "esesja_per_page":
        # Inny adapter, np. assemblies/kujawsko-pomorskie/
        raise NotImplementedError("Use sejmik-specific adapter")
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# Konstanty: kategorie wyników głosowania w eSesja
# ---------------------------------------------------------------------------

# Pattern, klucz wewnętrzny. Pattern matchuje header kategorii w "Wyniki
# imienne:" sekcji, np. "WSTRZYMUJĘ SIĘ" lub "WSTRZYMAŁ SIĘ" (oba spotykane
# zależnie od konfiguracji eSesja).
CATEGORIES: list[tuple[str, str]] = [
    (r"ZA", "za"),
    (r"PRZECIW", "przeciw"),
    (r"WSTRZYM\S+\s+SI[EĘ]", "wstrzymal_sie"),
    (r"BRAK G[ŁL]OSU", "brak_glosu"),
    (r"NIEOBECNI", "nieobecni"),
]

# Counts header w sumarycznym wierszu: "ZA: N, PRZECIW: N, ..."
COUNTS_RE = re.compile(
    r"ZA:\s*(\d+),\s*PRZECIW:\s*(\d+),"
    r"\s*WSTRZYM[^\d:,]+:\s*(\d+),"
    r"\s*BRAK G[ŁL]OSU:\s*(\d+),"
    r"\s*NIEOBECNI:\s*(\d+)"
)

# Footer eSesja, usuwany przy joinie tekstu z wielu stron
ESESJA_FOOTER_RE = re.compile(
    r"Wygenerowano za pomo[cć][aą] app\.esesja\.pl\s*"
    r"(?:\n?\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2})?"
)

# Czysty timestamp na końcu strony (po footerze)
BARE_TIMESTAMP_RE = re.compile(r"\n\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}\s*\n")

# Polskie miesiące do parsowania dat
POLISH_MONTHS = {
    "stycznia": "01", "lutego": "02", "marca": "03", "kwietnia": "04",
    "maja": "05", "czerwca": "06", "lipca": "07", "sierpnia": "08",
    "września": "09", "października": "10", "listopada": "11", "grudnia": "12",
}

# Rzymski -> arabski (dla numerów sesji)
ROMAN_TO_ARABIC = {
    "I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7, "VIII": 8,
    "IX": 9, "X": 10, "XI": 11, "XII": 12, "XIII": 13, "XIV": 14, "XV": 15,
    "XVI": 16, "XVII": 17, "XVIII": 18, "XIX": 19, "XX": 20, "XXI": 21,
    "XXII": 22, "XXIII": 23, "XXIV": 24, "XXV": 25, "XXVI": 26, "XXVII": 27,
    "XXVIII": 28, "XXIX": 29, "XXX": 30, "XXXI": 31, "XXXII": 32, "XXXIII": 33,
    "XXXIV": 34, "XXXV": 35, "XXXVI": 36, "XXXVII": 37, "XXXVIII": 38,
    "XXXIX": 39, "XL": 40, "XLI": 41, "XLII": 42, "XLIII": 43, "XLIV": 44,
    "XLV": 45, "XLVI": 46, "XLVII": 47, "XLVIII": 48, "XLIX": 49, "L": 50,
}


# ---------------------------------------------------------------------------
# Niskopoziomowe utility
# ---------------------------------------------------------------------------


def extract_pdf_text(pdf_path: str | Path, *, ocr_fallback: bool = True) -> tuple[str, str]:
    """Wyciąga tekst z PDF używając pdfplumber.

    Args:
        pdf_path: ścieżka do PDF
        ocr_fallback: jeśli pdfplumber zwróci 0 chars (skanowany PDF),
            spróbuj OCR przez pytesseract. Wymaga zainstalowanego
            Tesseract binary + pytesseract + pdf2image + Poppler.

    Returns:
        (full_text, first_page_text) – pełna treść i sam pierwsza strona.

    Cleanup: usuwa eSesja footer i "bare timestamps".
    """
    import pdfplumber

    with pdfplumber.open(str(pdf_path)) as pdf:
        full_text = ""
        first_page_text = ""
        total_chars = 0
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            total_chars += len(text.strip())
            if i == 0:
                first_page_text = text
            full_text += text + "\n"

    # Jeśli PDF jest skanowany (brak chars), pdfplumber zwraca pustki.
    # OCR fallback przez tesseract + pdf2image.
    if total_chars < 50 and ocr_fallback:
        full_text, first_page_text = _extract_with_ocr(pdf_path)

    full_text = ESESJA_FOOTER_RE.sub("", full_text)
    full_text = BARE_TIMESTAMP_RE.sub("\n", full_text)
    return full_text, first_page_text


def _extract_with_ocr(pdf_path: str | Path) -> tuple[str, str]:
    """OCR dla skanowanych PDFów (kujawsko-pomorskie, zachodniopomorskie itp).

    Wymaga: pytesseract + pdf2image + Poppler binary + Tesseract binary
    z polskim language pack (`tesseract-ocr-pol`).

    Raises:
        ImportError jeśli brakuje pytesseract/pdf2image
        FileNotFoundError jeśli Tesseract/Poppler nie zainstalowane
    """
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError as e:
        raise ImportError(
            "OCR fallback wymaga pytesseract + pdf2image. "
            "Zainstaluj: pip install pytesseract pdf2image, "
            "plus binaries: apt install tesseract-ocr tesseract-ocr-pol poppler-utils"
        ) from e

    images = convert_from_path(str(pdf_path), dpi=200)
    full_text = ""
    first_page_text = ""
    for i, img in enumerate(images):
        text = pytesseract.image_to_string(img, lang="pol") or ""
        if i == 0:
            first_page_text = text
        full_text += text + "\n"
    return full_text, first_page_text


def parse_polish_date(text: str) -> str | None:
    """Próbuje znaleźć datę sesji w tekście w 3 typowych formatach.

    Returns:
        ISO date "YYYY-MM-DD" jeśli znaleziono, inaczej None.
    """
    # Format DD.MM.YYYY: "Dnia 13.12.2024r." albo "z dnia: 27.04.2026"
    m = re.search(r"[Dd]nia\s*:?\s*(\d{1,2})\.(\d{1,2})\.(\d{4})", text)
    if m:
        d, mo, y = m.groups()
        return f"{y}-{int(mo):02d}-{int(d):02d}"

    # Format "w dniu DD miesiąc YYYY" (Polish months)
    month_pat = "|".join(POLISH_MONTHS.keys())
    m = re.search(rf"w dniu\s+(\d{{1,2}})\s+({month_pat})\s+(\d{{4}})", text)
    if m:
        d, mname, y = m.groups()
        return f"{y}-{POLISH_MONTHS[mname]}-{int(d):02d}"

    # Format "DD miesiąc YYYY" gdziekolwiek w tekście (najbardziej greedy)
    m = re.search(rf"(\d{{1,2}})\s+({month_pat})\s+(\d{{4}})", text)
    if m:
        d, mname, y = m.groups()
        return f"{y}-{POLISH_MONTHS[mname]}-{int(d):02d}"

    return None


def parse_session_number(first_page_text: str, filename: str = "") -> tuple[str | None, int | None]:
    """Wyciąga numer sesji w formie rzymskiej i arabskiej.

    Source priority:
    1. Treść pierwszej strony PDF (typowe: świętokrzyskie, gdański format)
    2. Nazwa pliku PDF (typowe: lubelski format `raport_z_glosowan_-_xxi_sesja_...`)

    Returns:
        (roman, arabic). Jedno albo oba mogą być None.
    """
    # "XXI sesja", "XXI Sesja Sejmiku", "Sesja XXI" w treści
    for pattern in (
        r"\b([IVXLCDM]+)\s+[Ss]esj",
        r"[Ss]esj\w*\s+([IVXLCDM]+)",
        r"^([IVXLCDM]+)\s+Protokół",
    ):
        m = re.search(pattern, first_page_text, re.MULTILINE)
        if m:
            roman = m.group(1)
            arabic = ROMAN_TO_ARABIC.get(roman)
            return roman, arabic

    # Arabic w treści: "Sesja nr 12" albo "12 Sesja"
    m = re.search(r"[Ss]esja\s+(?:nr\s+)?(\d+)", first_page_text)
    if m:
        arabic = int(m.group(1))
        arabic_to_roman = {v: k for k, v in ROMAN_TO_ARABIC.items()}
        return arabic_to_roman.get(arabic), arabic

    # Fallback: parse z filename. Pattern: "..._xxi_sesja_..." lub "..._XXI_..."
    if filename:
        fname_lower = filename.lower()
        # Match roman numeral surrounded by underscores or dashes
        m = re.search(r"[_\-]([ivxlcdm]+)[_\-]sesj", fname_lower)
        if m:
            roman = m.group(1).upper()
            arabic = ROMAN_TO_ARABIC.get(roman)
            if arabic:
                return roman, arabic
        # Match "sesja_N" arabic
        m = re.search(r"sesj\w*[_\-](\d+)", fname_lower)
        if m:
            arabic = int(m.group(1))
            arabic_to_roman = {v: k for k, v in ROMAN_TO_ARABIC.items()}
            return arabic_to_roman.get(arabic), arabic

    return None, None


# ---------------------------------------------------------------------------
# Parser głównego content: lista głosowań
# ---------------------------------------------------------------------------


def _parse_vote_block(
    block: str, vote_idx: int, *, vote_type: str = "uchwala"
) -> dict[str, Any] | None:
    """Parsuje pojedynczy blok głosowania (między dwoma "Głosowano w sprawie:").

    Args:
        block: tekst od końca "Głosowano w sprawie:" do następnego głosowania
        vote_idx: numer kolejny głosowania w sesji (0-based)
        vote_type: "uchwala" albo "wniosek" (na podstawie prefiksu)

    Returns: dict z `vote_index`, `vote_type`, `topic`, opcjonalnie `druk`,
    `resolution`, `counts`, `named_votes`, `voted_at`. None gdy blok nie
    wygląda jak głosowanie.
    """
    vote: dict[str, Any] = {"vote_index": vote_idx, "vote_type": vote_type}

    # Topic: od początku bloku do najwcześniejszego terminatora.
    # PDF eSesja ma "Wyniki głosowania" jako separator topic/wyniki.
    # DOCX małopolskie ma "Wyniki głosowania" przed "Głosowano w sprawie",
    # więc w bloku jest tylko sumarycznym "ZA: N, PRZECIW: ..." jako koniec topica.
    topic_match = re.match(
        r"(.*?)(?:Wyniki głosowania|ZA:\s*\d+,\s*PRZECIW)",
        block,
        re.DOTALL,
    )
    if topic_match:
        topic = re.sub(r"\s+", " ", topic_match.group(1)).strip()
        # Niektóre PDFy (lubelskie) mają "- czas głosowania: ..." w topicu.
        # Wyciągnij i przenieś do voted_at.
        czas_match = re.search(
            r"\s*-?\s*czas g[łl]osowania:\s*(.+?)$",
            topic,
            re.IGNORECASE,
        )
        if czas_match:
            vote["voted_at_raw"] = czas_match.group(1).strip()
            topic = topic[: czas_match.start()].strip().rstrip("-").strip()
        vote["topic"] = topic

    # Druk: "(druk N)"
    druk_match = re.search(r"\(druk\s+(\d+)\)", vote.get("topic", ""))
    if druk_match:
        vote["druk"] = int(druk_match.group(1))

    # Numer uchwały: "XXIII/562/26" w pierwszych 500 znakach bloku
    reso_match = re.search(r"([IVXLCDM]+/\d+/\d+)", block[:500])
    if reso_match:
        vote["resolution"] = reso_match.group(1)

    # Counts: "ZA: N, PRZECIW: N, ..."
    counts_match = COUNTS_RE.search(block)
    if counts_match:
        vote["counts"] = {
            "za": int(counts_match.group(1)),
            "przeciw": int(counts_match.group(2)),
            "wstrzymal_sie": int(counts_match.group(3)),
            "brak_glosu": int(counts_match.group(4)),
            "nieobecni": int(counts_match.group(5)),
        }

    # Data głosowania: "Głosowanie z dnia: DD.MM.YYYY, HH:MM:SS"
    # (świętokrzyskie format, eSesja standardowy footer)
    dnia_match = re.search(
        r"G[łl]osowani[ea] z dnia:?\s*(\d{1,2}\.\d{1,2}\.\d{4})(?:[,\s]+(\d{1,2}:\d{2}(?::\d{2})?))?",
        block,
    )
    if dnia_match:
        date_str = dnia_match.group(1)
        time_str = dnia_match.group(2) or ""
        d, m, y = date_str.split(".")
        iso = f"{y}-{int(m):02d}-{int(d):02d}"
        vote["voted_at"] = f"{iso} {time_str}" if time_str else iso

    # Named votes: split na "Wyniki imienne:"
    named_parts = re.split(r"Wyniki imienne:?\s*\n", block)
    if len(named_parts) > 1:
        vote["named_votes"] = _parse_named_section(named_parts[1])

    if vote.get("topic"):
        return vote
    return None


def _parse_named_section(named_text: str) -> dict[str, list[str]]:
    """Parsuje sekcję "Wyniki imienne" – nazwiska per kategoria głosu.

    Format:
        ZA (N)
        Imię Nazwisko, Imię Nazwisko, ...
        PRZECIW (N)
        Imię Nazwisko, ...
        WSTRZYMUJĘ SIĘ (N)
        ...
    """
    result: dict[str, list[str]] = {}

    # Find category positions
    cat_positions = []
    seen_keys = set()
    for pattern, key in CATEGORIES:
        m = re.search(pattern + r"\s*\(\d+\)", named_text)
        if m and key not in seen_keys:
            cat_positions.append((m.start(), m.end(), pattern, key))
            seen_keys.add(key)

    cat_positions.sort(key=lambda x: x[0])

    for idx_c, (start, _hdr_end, pat, key) in enumerate(cat_positions):
        # Names start after "LABEL (N)\n"
        hdr_match = re.search(pat + r"\s*\(\d+\)\s*\n", named_text[start:])
        if not hdr_match:
            result[key] = []
            continue

        names_start = start + hdr_match.end()

        # Names end at next category start or at end markers
        if idx_c < len(cat_positions) - 1:
            names_end = cat_positions[idx_c + 1][0]
        else:
            rest = named_text[names_start:]
            end_match = re.search(
                r"\n\d+[\.\)]\s+|Głosowano w sprawie|Głosowanie z dnia|Wygenerowano za pomocą",
                rest,
            )
            names_end = names_start + end_match.start() if end_match else len(named_text)

        names_text = named_text[names_start:names_end].strip()
        names_text = ESESJA_FOOTER_RE.sub("", names_text)
        names_text = re.sub(r"\s+", " ", names_text).strip()

        if not names_text:
            result[key] = []
            continue

        result[key] = _filter_valid_names(
            n.strip() for n in names_text.split(",") if n.strip()
        )

    # Categories not present in text = empty list
    present = {k for _, _, _, k in cat_positions}
    for _pat, key in CATEGORIES:
        if key not in present:
            result[key] = []

    return result


def _filter_valid_names(raw_names) -> list[str]:
    """Filtr na heurystycznie poprawne nazwiska radnych.

    Reguły:
    - 2 do 4 słów
    - długość ≤35 znaków
    - zaczyna się dużą literą
    - brak cyfr
    - brak znaków specjalnych ();./

    Wyłącza śmieci typu "Strona 1", timestampy, fragmenty page-break itp.
    """
    result = []
    for n in raw_names:
        words = n.split()
        if (
            2 <= len(words) <= 4
            and len(n) <= 35
            and n[0].isupper()
            and not re.search(r"\d", n)
            and not any(c in n for c in "();./")
        ):
            result.append(n)
    return result


def extract_docx_text(docx_path: str | Path) -> str:
    """Wyciąga tekst z DOCX. Używane dla małopolskiego, gdzie głosowania są
    publikowane jako osobne docxy (1 docx = 1 głosowanie).

    Format treści docx jest identyczny z eSesja PDF: "Głosowano w sprawie / ...
    / Wyniki imienne / ZA (N) nazwiska...". Dlatego można reuse parser.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "DOCX parsing wymaga python-docx. Zainstaluj: pip install python-docx"
        ) from e

    doc = Document(str(docx_path))
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # Tabele w docx (małopolskie nie ma, ale na wypadek)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Detekcja formatu PDF
# ---------------------------------------------------------------------------


def detect_pdf_format(pdf_path: str | Path) -> str:
    """Wykrywa format PDF z imiennymi głosowaniami żeby wskazać który parser użyć.

    Returns: jedna z wartości:
        "esesja_standard" - tabela "Głosowano w sprawie / Wyniki imienne",
            parse_voting_pdf() w tej lib.
        "esesja_per_page" - 1 głosowanie per strona z tabelą Lp/Nazwisko/Decyzja,
            wymaga osobnego parsera (nie w tej lib).
        "scanned" - skanowany PDF (chars=0), wymaga OCR + ponowna detekcja.
        "narrative" - protokół narracyjny (proza), nie ma tabel głosowań.
        "unknown" - format nie rozpoznany.
    """
    import pdfplumber

    with pdfplumber.open(str(pdf_path)) as pdf:
        # Sumuj chars z 3 pierwszych stron
        total_chars = 0
        sample_text = ""
        for page in pdf.pages[:3]:
            text = page.extract_text() or ""
            total_chars += len(text.strip())
            sample_text += text + "\n"

    if total_chars < 50:
        return "scanned"

    # Sygnatury formatu eSesja_standard
    if (
        re.search(r"G[łl]osowano(?:\s+wniosek)?\s+w sprawie:", sample_text)
        and "Wyniki imienne" in sample_text
    ):
        return "esesja_standard"

    # Sygnatury formatu eSesja_per_page (kuj-pom, podobne)
    if (
        re.search(r"Typ\s+g[łl]osowania", sample_text, re.IGNORECASE)
        and re.search(r"Data\s+g[łl]osowania", sample_text, re.IGNORECASE)
    ):
        return "esesja_per_page"

    # Sygnatury narracyjnego protokołu (case insensitive bo czasem ALL CAPS)
    if re.search(r"Protok[óo][łl]\s+(?:nr\s+)?[IVXLCDM]+", sample_text, re.IGNORECASE):
        # Sprawdź czy oprócz "Protokół nr X" są też tabele głosowań
        if "Wyniki imienne" not in sample_text and "Głosowano" not in sample_text:
            return "narrative"

    return "unknown"


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def parse_voting_text(
    full_text: str, first_page: str = "", source_name: str = ""
) -> dict[str, Any]:
    """Parsuje już wyciągnięty tekst eSesja (z PDF, DOCX, OCR czy gdziekolwiek).

    Najniższy poziom API. Używane jako wewnętrzny core przez
    parse_voting_pdf() i parse_voting_docx().

    Args:
        full_text: cały tekst dokumentu po cleanup (usunięte eSesja footers)
        first_page: tekst pierwszej strony (dla metadanych typu data, numer sesji)
        source_name: nazwa źródła (filename) do fallback parsowania numeru sesji

    Returns: dict z metadanymi sesji + listą głosowań.
    """
    iso_date = parse_polish_date(first_page or full_text[:2000]) or parse_polish_date(full_text)
    roman, arabic = parse_session_number(first_page or full_text[:2000], filename=source_name)

    session: dict[str, Any] = {
        "source_file": source_name,
        "date": iso_date,
        "number_roman": roman,
        "number": arabic,
    }

    # Split na bloki głosowań
    parts = re.split(
        r"G[łl]osowano(\s+wniosek)?\s+w sprawie:\s*",
        full_text,
    )
    votes = []
    for idx in range(1, len(parts), 2):
        wniosek_marker = parts[idx]
        block = parts[idx + 1] if idx + 1 < len(parts) else ""
        vote_type = "wniosek" if wniosek_marker else "uchwala"
        vote = _parse_vote_block(block, len(votes), vote_type=vote_type)
        if vote:
            votes.append(vote)

    session["votes"] = votes
    session["vote_count"] = len(votes)
    return session


def parse_voting_pdf(pdf_path: str | Path) -> dict[str, Any]:
    """Parsuje PDF eSesja z imiennymi wykazami głosowań do struktury JSON.

    Returns:
        dict z polami:
            source_file: str — basename PDF
            date: str | None — ISO date YYYY-MM-DD
            number_roman: str | None — np. "XXI"
            number: int | None — np. 21
            votes: list[dict] — lista głosowań
            vote_count: int

        Każde głosowanie w `votes`:
            vote_index: int — numer kolejny 0-based
            vote_type: "uchwala" | "wniosek"
            topic: str — tytuł uchwały
            druk: int (optional) — numer druku
            resolution: str (optional) — numer uchwały
            counts: {za, przeciw, wstrzymal_sie, brak_glosu, nieobecni}
            named_votes: {za: [...], przeciw: [...], ...}
            voted_at: str (optional) — ISO datetime z "Głosowanie z dnia: ..."
    """
    full_text, first_page = extract_pdf_text(pdf_path)
    return parse_voting_text(full_text, first_page, source_name=Path(pdf_path).name)


# ---------------------------------------------------------------------------
# Per-page format: 1 strona PDF = 1 głosowanie z tabelą Lp/Nazwisko/Decyzja
# (kujawsko-pomorskie, lubuskie, zachodniopomorskie - skanowane PDF + OCR)
# ---------------------------------------------------------------------------


# Decyzja w OCR (wartości w trzeciej kolumnie tabeli):
PER_PAGE_DECYZJA_MAP = {
    "ZA": "za",
    "PRZECIW": "przeciw",
    "WSTRZYMAŁ SIĘ": "wstrzymal_sie",
    "WSTRZYMAL SIE": "wstrzymal_sie",
    "WSTRZYMUJĘ SIĘ": "wstrzymal_sie",
    "WSTRZYMUJE SIE": "wstrzymal_sie",
    "NIE GŁOSOWAŁ": "brak_glosu",
    "NIE GLOSOWAL": "brak_glosu",
    "NIEOBECNY": "nieobecni",
    "NIEOBECNA": "nieobecni",
    "BRAK": "brak_glosu",
    "BRAK GŁOSU": "brak_glosu",
}


def _parse_per_page_vote(page_text: str, page_idx: int) -> dict[str, Any] | None:
    """Parsuje 1 stronę PDF (kuj-pom, lubuskie, zach-pom format).

    Heurystyki na tekst po OCR (działa nawet z polish characters OCR'd jako
    łacińskie alternatywy - "Głosowanie" → "Gtosowanie", "ć" → "c" itp):
    1. Header "Typ głosowania: jawne   Data głosowania: DD.MM.YYYY HH:MM"
    2. Topic - heurystycznie z early text strony, ale bywa zniekształcony
    3. Tabela z liniami "Lp Nazwisko Decyzja". Dwie kolumny, więc czasem 2 wpisy
       w jednej linii. Decyzje (ZA/NIEOBECNY itp) są przeważnie OCR'd dobrze
       bo są short caps.

    Detect: strona jest "głosowaniem" jeśli zawiera ≥5 wzorców "nazwisko + decyzja"
    (nie wymagamy słowa "Głosowanie" bo OCR może je zniekształcić).

    Returns: vote dict albo None gdy strona nie wygląda jak głosowanie.
    """
    vote: dict[str, Any] = {
        "vote_index": page_idx,
        "vote_type": "uchwala",
    }

    # Data głosowania
    date_match = re.search(
        r"[Dd]ata\s+g[łl]osowania:?\s*(\d{1,2}\.\d{1,2}\.\d{4})"
        r"(?:[\s,]+(\d{1,2}:\d{2}(?::\d{2})?))?",
        page_text,
        re.IGNORECASE,
    )
    if date_match:
        d, m, y = date_match.group(1).split(".")
        iso = f"{y}-{int(m):02d}-{int(d):02d}"
        vote["voted_at"] = f"{iso} {date_match.group(2)}" if date_match.group(2) else iso

    # Topic - heurystycznie z pierwszych 600 znaków przed "Typ głosowania"
    # albo "Liczba uprawnionych" (start tabeli wyników). OCR'd topic może
    # być uszkodzony, ale ma wskaźnik tematu uchwały.
    topic_match = re.match(
        r"(.+?)(?:Typ\s+g[łl]?osowania|[LtT]iczba\s+uprawnionych)",
        page_text[:1500],
        re.IGNORECASE | re.DOTALL,
    )
    if topic_match:
        topic = re.sub(r"\s+", " ", topic_match.group(1)).strip()
        # Usuwamy nagłówek sesji jeśli na początku
        topic = re.sub(r"^.*?[Ss]esj\S*\s+[Ss]ejm\S*\s+\S+\s+\S+", "", topic).strip()
        vote["topic"] = topic[:300] if topic else ""

    # Tabela: linie zaczynające się od "N." lub "N ", z nazwiskiem i decyzją na końcu
    # Może być 2-kolumnowa, więc trzeba potraktować dwukrotnie.
    named: dict[str, list[str]] = {key: [] for key in PER_PAGE_DECYZJA_MAP.values()}

    # Pattern który łapie: numer + nazwisko + decyzja
    # OCR często ma artefakty, więc tolerujemy spacje i niewłaściwe litery
    # Decyzja musi być całym słowem na końcu / w środku linii
    decyzja_alt = "|".join(re.escape(k) for k in PER_PAGE_DECYZJA_MAP)
    row_re = re.compile(
        rf"(?:\d+[.,]?\s*)?([A-ZĄĆĘŁŃÓŚŹŻ][\wÀ-ſ \-]{{2,40}}?)\s+({decyzja_alt})\b",
    )

    for m in row_re.finditer(page_text):
        name_raw, decision = m.groups()
        name = re.sub(r"\s+", " ", name_raw).strip(" .,;")
        # Filtr na właściwe nazwiska
        words = name.split()
        if len(words) < 2 or len(words) > 4:
            continue
        if len(name) > 35 or len(name) < 5:
            continue
        # Skip "Lp Nazwisko" header itp.
        if any(w in {"Nazwisko", "Imie", "Imię", "Pp", "Lp", "Typ", "Liczba"}
               for w in words):
            continue
        key = PER_PAGE_DECYZJA_MAP[decision]
        if name not in named[key]:
            named[key].append(name)

    vote["named_votes"] = named
    vote["counts"] = {k: len(v) for k, v in named.items()}

    # Wymagamy żeby strona miała przynajmniej kilka rozpoznanych głosów,
    # inaczej to nie była tabela głosowań (np. okładka, podsumowanie).
    total_names = sum(vote["counts"].values())
    if total_names < 5:
        return None

    return vote


def parse_voting_pdf_per_page(pdf_path: str | Path) -> dict[str, Any]:
    """Parser dla skanowanych PDF z formatem "1 strona = 1 głosowanie".

    Dotyczy: kujawsko-pomorskie, lubuskie, zachodniopomorskie. Każda strona
    PDF to osobne głosowanie z tabelą Lp/Nazwisko/Decyzja. Wymaga OCR
    (Tesseract z polskim pack dla najlepszej accuracy; eng pack daje ~85%).

    Returns: tak samo jak parse_voting_pdf().
    """
    try:
        import pdfplumber
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError as e:
        raise ImportError(
            "parse_voting_pdf_per_page wymaga pdf2image + pytesseract. "
            "Plus binaries: tesseract-ocr-pol (apt) lub eng pack jako fallback."
        ) from e

    # Najpierw spróbuj wyciągnąć tekst standardowo (na wypadek gdyby PDF
    # nie był skanem). Jeśli puste, OCR per strona.
    pages_text: list[str] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            pages_text.append(txt)

    total_chars = sum(len(t.strip()) for t in pages_text)
    if total_chars < 50 * len(pages_text):
        # Skanowany PDF, OCR każdą stronę
        # Spróbuj polski pack, fallback do eng
        lang = "pol"
        try:
            pytesseract.get_languages(config="")
            available = pytesseract.get_languages()
            if "pol" not in available:
                lang = "eng"
        except Exception:
            lang = "eng"

        images = convert_from_path(str(pdf_path), dpi=250)
        pages_text = []
        for img in images:
            pages_text.append(pytesseract.image_to_string(img, lang=lang))

    # Per strona = 1 głosowanie
    votes = []
    for idx, page_text in enumerate(pages_text):
        vote = _parse_per_page_vote(page_text, len(votes))
        if vote:
            votes.append(vote)

    # Metadata z pierwszej strony
    first_page = pages_text[0] if pages_text else ""
    full_text = "\n".join(pages_text)
    iso_date = parse_polish_date(first_page) or parse_polish_date(full_text)
    roman, arabic = parse_session_number(first_page, filename=Path(pdf_path).name)

    return {
        "source_file": Path(pdf_path).name,
        "date": iso_date,
        "number_roman": roman,
        "number": arabic,
        "votes": votes,
        "vote_count": len(votes),
    }


def parse_voting_docx(docx_path: str | Path) -> dict[str, Any]:
    """Parsuje DOCX eSesja (małopolskie format). 1 docx = 1 głosowanie.

    Treść docx jest identyczna z formatem eSesja PDF, tylko zapakowana
    w osobnym pliku Word per uchwała. Reuse parse_voting_text() po wyciągnięciu
    tekstu z docx przez python-docx.

    Returns: tak samo jak parse_voting_pdf() ale `votes` typowo ma 1 element.
    """
    full_text = extract_docx_text(docx_path)
    return parse_voting_text(full_text, "", source_name=Path(docx_path).name)


def validate_parsed(session: dict[str, Any]) -> tuple[int, int, list[str]]:
    """Walidacja: czy liczba nazwisk w named_votes zgadza się z counts.

    Returns: (ok_count, fail_count, errors).

    Toleruje brak counts header (rekonstruuje z named_votes) i drobne mismatches.
    """
    errors = []
    ok = 0
    keys = ["za", "przeciw", "wstrzymal_sie", "brak_glosu", "nieobecni"]
    for i, v in enumerate(session["votes"]):
        c = v.get("counts", {})
        nv = v.get("named_votes", {})
        expected = sum(c.get(k, 0) for k in keys)
        actual = sum(len(nv.get(k, [])) for k in keys)
        if expected == actual:
            ok += 1
        elif expected == 0 and actual > 0:
            # Counts not parsed but names present – reconstruct
            for k in keys:
                c[k] = len(nv.get(k, []))
            v["counts"] = c
            v["counts_reconstructed"] = True
            ok += 1
        else:
            diff = actual - expected
            sign = "+" if diff > 0 else ""
            errors.append(f"Vote {i}: expected {expected}, got {actual} ({sign}{diff})")
    return ok, len(errors), errors


# ---------------------------------------------------------------------------
# CLI (do testów)
# ---------------------------------------------------------------------------


if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 2:
        print("Usage: lib_voting_pdf_table.py <pdf_file> [...]", file=sys.stderr)
        sys.exit(1)

    for path in sys.argv[1:]:
        result = parse_voting_pdf(path)
        ok, fail, errors = validate_parsed(result)
        print(f"\n=== {path} ===")
        print(f"  date: {result['date']}, sesja: {result['number_roman']}/{result['number']}")
        print(f"  głosowań: {result['vote_count']}")
        print(f"  walidacja: {ok} OK / {fail} FAIL")
        for e in errors[:3]:
            print(f"    - {e}")
        # Sample pierwsze głosowanie
        if result["votes"]:
            v = result["votes"][0]
            print(f"  Sample vote 0:")
            print(f"    topic: {v.get('topic', '?')[:80]}")
            print(f"    counts: {v.get('counts')}")
            for cat in ("za", "przeciw", "wstrzymal_sie", "brak_glosu", "nieobecni"):
                nv = v.get("named_votes", {}).get(cat, [])
                if nv:
                    print(f"    {cat} ({len(nv)}): {nv[:3]}{'...' if len(nv) > 3 else ''}")
