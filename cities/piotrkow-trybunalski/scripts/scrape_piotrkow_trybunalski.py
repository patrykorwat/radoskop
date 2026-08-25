#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Radoskop Piotrków Trybunalski — imienne głosowania Rady Miasta.

Źródło: BIP UM Piotrkowa Trybunalskiego (www.bip.piotrkow.pl, platforma AkcessNet).
Rada Miasta (IX kadencja 2024-2029) publikuje w kategoriach
"Wyniki głosowań z sesji Rady Miasta - {rok}" per-sesyjne artykuły, z jednym
załącznikiem DOCX na głosowanie. Każdy DOCX to tabela imienna:
    Typ głosowania … | Data głosowania: DD.MM.YYYY HH:MM
    Uprawnionych: N | Za: … / Przeciw / Wstrzymało się / Nieobecni
    Lp. | Imię i nazwisko | Głos | Data i czas oddania głosu
    1 | Jan Dziemdziora | Za | 24.06.2026 11:12
Zapisuje standardowe formaty Radoskopa: kadencja-2024-2029.json + profiles.json
+ data.json (indeks).

Głosowania "tajne" (np. wybór Przewodniczącego/komisji) nie mają tabeli imiennej
-> są pomijane (nie ma danych per radny). Załączniki nie-DOCX (np. wyjaśnienia
radnych, skany imiennych wykazów z sesji nadzwyczajnych) pomijane.

Użycie:
    python scrape_piotrkow_trybunalski.py --output docs/data.json \
        --profiles docs/profiles.json [--cache-dir .cache]
"""

import argparse
import hashlib
import io
import json
import re
import sys
import time
import unicodedata
from collections import Counter, defaultdict
from datetime import datetime
from itertools import combinations
from pathlib import Path

import requests
from docx import Document

BIP = "https://www.bip.piotrkow.pl"
YEARS = [  # (label, catid, x) — kategorie "Wyniki głosowań ... - {rok}"
    ("2024", "2310", 50),
    ("2025", "2387", 49),
    ("2026", "2474", 48),
]
KAD_START = "2024-05-07"
KADENCJA_ID = "2024-2029"
KADENCJA_LABEL = "IX kadencja (2024\u20132029)"
PAGE_CAP = 8
REQ_DELAY = 0.6
_LAST_REQ = 0.0

UA = {"User-Agent": "Mozilla/5.0 (X11; Linux x86_64) Radoskop/1.0"}

# Kluby radnych — PENDING (do kuratorowania z BIP; nie fabrykować).
CLUB_ASSIGN: dict = {}
_CLUBS_META = {}


def _rate():
    global _LAST_REQ
    now = time.time()
    d = now - _LAST_REQ
    if d < REQ_DELAY:
        time.sleep(REQ_DELAY - d)
    _LAST_REQ = time.time()


def fetch(url, cache_dir=None, binary=False, tries=6):
    if cache_dir is not None:
        key = hashlib.md5(url.encode()).hexdigest()
        ext = ".bin" if binary else ".html"
        cf = cache_dir / (key + ext)
        if cf.is_file():
            return cf.read_bytes() if binary else cf.read_text(encoding="utf-8", errors="ignore")
    for i in range(tries):
        _rate()
        try:
            resp = requests.get(url, headers=UA, timeout=45)
            resp.raise_for_status()
            data = resp.content if binary else resp.text
            if not binary and len(data) < 5000 and "/upload/plik," not in data:
                time.sleep(2 + i)
                continue
            if cache_dir is not None:
                cf = cache_dir / (key + ext)
                cf.parent.mkdir(parents=True, exist_ok=True)
                if binary:
                    cf.write_bytes(data)
                else:
                    cf.write_text(data, encoding="utf-8", errors="ignore")
            return data
        except Exception:
            time.sleep(2 + i * 2)
    raise RuntimeError(f"fetch failed: {url}")


# ---- Normalizacja nazwisk / slugi -----------------------------------------
def _norm(s):
    s = s.lower().replace("\u0142", "l").replace("\u0141", "L")
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    return re.sub(r"[^a-z0-9]+", "", s)


def make_slug(name):
    repl = {'\u0105': 'a', '\u0107': 'c', '\u0119': 'e', '\u0142': 'l', '\u0144': 'n',
            '\u00f3': 'o', '\u015b': 's', '\u017a': 'z', '\u017c': 'z',
            '\u0104': 'A', '\u0106': 'C', '\u0118': 'E', '\u0141': 'L', '\u0143': 'N',
            '\u00d3': 'O', '\u015a': 'S', '\u0179': 'Z', '\u017b': 'Z'}
    slug = name.lower()
    for pl, a in repl.items():
        slug = slug.replace(pl, a)
    return re.sub(r"[^a-z0-9]+", "", slug)


def _club_of(name):
    return CLUB_ASSIGN.get(name, "NZ")


# Warianty pisowni nazwisk w źródłowych DOCX (ten sam radny) -> kanon.
_CANON = {
    "Rafa\u0142 Czajka": "Rafa\u0142 Pawe\u0142 Czajka",
    "Przemys\u0142aw Winiarski": "Przemys\u0142aw Winiarski",  # whitespace handled below
}


def _canonical(name):
    name = re.sub(r"\s+", " ", name).strip()  # ze źródła bywa "Przemys\u0142aw  Winiarski"
    return _CANON.get(name, name)


# ---- 1. Kolekcja sesji + załączników DOCX z kategorii Wyniki głosowań ----
_SESS_RE = re.compile(
    r"Wynik(?:i|u)?\s+g\u0142osowania\s+z\s+(.+?)Sesji\s+(?:Rady\s+Miasta\s+)?"
    r"z\s+(?:dn\.|dnia)\s+(\d{1,2})\s*(?:\.|)\s*\(?(\d{1,2}|stycznia|lutego|marca|kwietnia|"
    r"maja|czerwca|lipca|sierpnia|wrze\u015bnia|pa\u017adziernika|listopada|grudnia)\)?\s*(\d{4})",
    re.IGNORECASE,
)
_MONTHS = {"stycznia": 1, "lutego": 2, "marca": 3, "kwietnia": 4, "maja": 5,
           "czerwca": 6, "lipca": 7, "sierpnia": 8, "wrze\u015bnia": 9,
           "pa\u017adziernika": 10, "listopada": 11, "grudnia": 12}


def _parse_header_date(header):
    ns = " ".join(header.split())
    # "z dnia 28 listopada 2024 r." / "z dn. 24.06.2026 r." / "z dnia 6 maja 2024 r."
    m = re.search(r"z\s+(?:dn\.|dnia)\s+(.+)", ns)
    if not m:
        return None
    frag = re.sub(r"\s+r\.?$", "", m.group(1)).strip()
    # numeric DD.MM.YYYY
    md = re.fullmatch(r"(\d{1,2})\.(\d{1,2})\.(\d{4})", frag)
    if md:
        return f"{md.group(3)}-{int(md.group(2)):02d}-{int(md.group(1)):02d}"
    # month-name: "6 maja 2024"
    mm = re.fullmatch(r"(\d{1,2})\s+(\w+)\s+(\d{4})", frag)
    if mm and mm.group(2).lower() in _MONTHS:
        return f"{mm.group(3)}-{_MONTHS[mm.group(2).lower()]:02d}-{int(mm.group(1)):02d}"
    return None


def _roman_to_int(roman):
    vals = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100, 'D': 500, 'M': 1000}
    n = 0
    prev = 0
    for ch in reversed(roman.upper()):
        v = vals.get(ch, 0)
        n += -v if v < prev else v
        prev = v
    return n


def _direct_url(href):
    """Zwraca bezpośredni URL uploadu (pomija docs.google viewer)."""
    if href.startswith(f"{BIP}/upload/"):
        return href
    m = re.search(r"url=({}[^&]+)".format(re.escape(f"{BIP}/upload/plik,")), href)
    return m.group(1) if m else None


def _parse_year(catid, x, cache_dir=None):
    """Iteruje strony kategorii i zwraca listę sesji z docx załącznikami."""
    sessions = []
    seen_sess = set()
    page = 0
    while page < PAGE_CAP:
        url = f"{BIP}/index.php?idg=15&id={catid}&x=19&y={x}"
        if page > 0:
            url += f"&a={page}"
        try:
            html = fetch(url, cache_dir)
        except Exception as e:
            print(f"    [warn] {catid} page {page}: {e}")
            break
        # sesje = bloki pub-box z nagłówkiem h2 i załącznikami
        heads = list(re.finditer(r'<h2[^>]*>(.*?)</h2>', html, re.S))
        if not heads:
            break
        n_added = 0
        for mi in range(len(heads)):
            start = heads[mi].start()
            end = heads[mi + 1].start() if mi + 1 < len(heads) else len(html)
            seg = html[start:end]
            htxt = re.sub(r"\s+", " ", re.sub("<[^>]+>", " ", heads[mi].group(1))).strip()
            if "glosow" not in htxt.lower() and "głosow" not in htxt.lower():
                continue
            date = _parse_header_date(htxt)
            # załączniki docx (pojedyncze, bez viewerów)
            atts = []
            for a in re.finditer(r'href="([^"]*upload/plik,[^"]*)"', seg):
                u = _direct_url(a.group(1))
                if not u or ".docx" not in u:
                    continue
                fid = u.split("/upload/plik,")[1].split(",")[0]
                if fid not in [x["fid"] for x in atts]:
                    atts.append({"fid": fid, "url": u})
            if not atts:
                continue
            roman = ""
            rm = re.search(r"z\s+(.+?)Sesji", htxt, re.I)
            roman = rm.group(1).replace("Nadzwyczajnej", "").replace("Rady Miasta", "").strip() if rm else ""
            key = (htxt, date)
            if key in seen_sess:
                continue
            seen_sess.add(key)
            sessions.append({"header": htxt, "date": date, "roman": roman,
                             "num": _roman_to_int(roman) if roman else 0,
                             "atts": atts})
            n_added += 1
        print(f"    {catid} page {page}: +{n_added} sessions (total {len(sessions)})")
        # paginacja: czy jest "następna strona"
        if "nast\u0119pna strona" in html:
            page += 1
            time.sleep(0.8)
        else:
            break
    return sessions


def collect_sessions(cache_dir=None):
    out = []
    for label, catid, x in YEARS:
        print(f"  -- rok {label} (cat {catid}) --")
        sesses = _parse_year(catid, x, cache_dir)
        for s in sesses:
            if s["date"] and s["date"] >= KAD_START:
                out.append(s)
        time.sleep(0.6)
    out.sort(key=lambda s: (s["date"] or "", s["num"]))
    return out


# ---- 2. Parsowanie DOCX głosowania ---------------------------------------
_VOTE_RE = re.compile(
    r"(Za|Przeciw|Wstrzyma\u0142 si\u0119|Wstrzyma\u0142a si\u0119|"
    r"Nie g\u0142osowa\u0142|Nie g\u0142osowa\u0142a|Nieobecny|Nieobecna|"
    r"NIE G\u0142OSOWA\u0141|WSTRZYMA\u0141 SI\u0118|PRZECIW|ZA)", re.IGNORECASE)


def _clean_vote(vote):
    v = vote.strip().lower()
    v = v.replace("\u0142", "l").replace("\u0119", "e")
    if v in ("za",):
        return "za"
    if "przeciw" in v:
        return "przeciw"
    if "wstrzym" in v:
        return "wstrzymal_sie"
    if "nie glosowal" in v or "nie głosował" in v:
        return "brak_glosu"
    if "nieobecn" in v:
        return "nieobecni"
    return None


def _docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_vote(data: bytes):
    """Parsuje jeden DOCX -> jeden rekord głosowania (lub None)."""
    try:
        txt = _docx_text(data)
    except Exception:
        return None
    # temat
    tm = re.search(r"G\u0142osowanie w sprawie:\s*(.+?)(?:\n|$)", txt, re.I)
    topic = tm.group(1).strip().rstrip(".").strip() if tm else ""
    if not topic:
        # nagłówek "5.1. Podjęcie uchwały..."
        m2 = re.search(r"^\s*\d+\.\d+\.\s+(.+)$", txt, re.M)
        if m2:
            topic = m2.group(1).strip().rstrip(".").strip()
    # data głosowania
    dm = re.search(r"Data g\u0142osowania:\s*(\d{1,2})\.(\d{1,2})\.(\d{4})", txt, re.I)
    vote_date = None
    if dm:
        vote_date = f"{dm.group(3)}-{int(dm.group(2)):02d}-{int(dm.group(1)):02d}"
    named = {"za": [], "przeciw": [], "wstrzymal_sie": [], "brak_glosu": [], "nieobecni": []}
    # tabele docx (wiersze "Lp | Name | Glos | timestamp")
    n_found = 0
    for line in txt.split("\n"):
        l = line.strip()
        if not l:
            continue
        row = re.match(
            r"^(\d{1,3})\s*\|\s*(.+?)\s*\|\s*(Za|Przeciw|Wstrzyma\u0142 si\u0119|"
            r"Wstrzyma\u0142a si\u0119|Nie g\u0142osowa\u0142|Nie g\u0142osowa\u0142a|"
            r"Nieobecny|Nieobecna)\s*\|", l, re.I)
        if row:
            name = _canonical(row.group(2))
            v = _clean_vote(row.group(3))
            if name and v:
                named[v].append(name)
                n_found += 1
    if n_found == 0:
        return None
    return {"session_date": vote_date, "topic": topic, "named": named,
            "counts": Counter(c for k, names in named.items() for c in ([k] * len(names)))}


# ---- 3. Główna kolekcja ----------------------------------------------------
def collect_all(sessions, cache_dir=None):
    records = []
    for s in sessions:
        nvotes = 0
        for att in s["atts"]:
            try:
                data = fetch(att["url"], cache_dir, binary=True)
            except Exception as e:
                print(f"    [warn] {s['roman']} {s['date']} pdf {att['fid']}: {e}")
                continue
            v = _parse_vote(data)
            if not v:
                continue
            rec = dict(v)
            rec["session_num"] = s["roman"]
            # data sesji: prefer głosowania z docx, fallback z nagłówka
            rec["session_date"] = v["session_date"] or s["date"]
            records.append(rec)
            nvotes += 1
        print(f"  {s['roman'] or '?':10s} {s['date']} votes={nvotes} (from {len(s['atts'])} docx)")
    return records


# ---- 4. Budowa wyjścia (port z scrape_konin.py) ---------------------------
def _canonical_rec(rec):
    for k in ("za", "przeciw", "wstrzymal_sie", "brak_glosu", "nieobecni"):
        rec["named"][k] = list(dict.fromkeys(rec["named"].get(k, [])))
    return rec


def _compute_consensus(all_votes):
    club_majority = {}
    for v in all_votes:
        by_club = defaultdict(list)
        for cat in ("za", "przeciw", "wstrzymal_sie"):
            for name in v["named_votes"].get(cat, []):
                by_club[_club_of(name)].append(cat)
        for cl, cats in by_club.items():
            if cats:
                club_majority[(cl, v["id"])] = Counter(cats).most_common(1)[0][0]
    stats = defaultdict(lambda: {"za": 0, "przeciw": 0, "wstrzymal": 0, "brak": 0,
                                 "nieobecny": 0, "with": 0, "against": 0, "sess": set()})
    for v in all_votes:
        for cat, names in v["named_votes"].items():
            for name in names:
                key = "za" if cat == "za" else "przeciw" if cat == "przeciw" \
                    else "wstrzymal" if cat == "wstrzymal_sie" \
                    else "nieobecny" if cat == "nieobecni" else "brak"
                stats[name][key] += 1
                if key != "nieobecny":
                    stats[name]["sess"].add(v["session_date"])
        for cat in ("za", "przeciw", "wstrzymal_sie"):
            for name in v["named_votes"].get(cat, []):
                maj = club_majority.get((_club_of(name), v["id"]))
                if maj is None:
                    continue
                if cat == maj:
                    stats[name]["with"] += 1
                else:
                    stats[name]["against"] += 1
    return club_majority, stats


def build_output(records):
    all_votes = []
    vid = 0
    sessions_by_date = {}
    for rec in records:
        d = rec.get("session_date")
        if not d:
            continue
        if d not in sessions_by_date:
            sessions_by_date[d] = {"date": d, "number": rec.get("session_num", ""),
                                   "vote_count": 0, "attendees": set()}
        named = {k: list(v) for k, v in rec["named"].items()}
        vid += 1
        sessions_by_date[d]["vote_count"] += 1
        for cat in ("za", "przeciw", "wstrzymal_sie", "brak_glosu", "nieobecni"):
            sessions_by_date[d]["attendees"].update(named.get(cat, []))
        all_votes.append({
            "id": str(vid), "session_date": d,
            "session_number": rec.get("session_num", ""),
            "topic": rec["topic"] or "", "named_votes": named,
            "counts": {k: len(named.get(k, [])) for k in ("za", "przeciw", "wstrzymal_sie")},
        })

    sessions_data = []
    for d in sorted(sessions_by_date.keys()):
        s = sessions_by_date[d]
        sessions_data.append({
            "date": d, "number": s["number"], "vote_count": s["vote_count"],
            "attendee_count": len(s["attendees"]), "attendees": sorted(s["attendees"]),
            "speakers": [],
        })

    all_names = set()
    for v in all_votes:
        for names in v["named_votes"].values():
            all_names.update(names)

    councilors_data = {}
    for name in sorted(all_names):
        councilors_data[name] = {
            "name": name, "club": _club_of(name), "district": None,
            "votes_za": 0, "votes_przeciw": 0, "votes_wstrzymal": 0,
            "votes_brak": 0, "votes_nieobecny": 0,
            "votes_with_club": 0, "votes_against_club": 0, "rebellions": [],
        }
    for v in all_votes:
        for cat, names in v["named_votes"].items():
            for name in names:
                if name not in councilors_data:
                    continue
                c = councilors_data[name]
                if cat == "za":
                    c["votes_za"] += 1
                elif cat == "przeciw":
                    c["votes_przeciw"] += 1
                elif cat == "wstrzymal_sie":
                    c["votes_wstrzymal"] += 1
                elif cat == "nieobecni":
                    c["votes_nieobecny"] += 1
                else:
                    c["votes_brak"] += 1

    total_votes = len(all_votes)
    total_sessions = len(sessions_data)

    _, stats = _compute_consensus(all_votes)

    councilors_list = []
    for name in sorted(councilors_data.keys()):
        c = councilors_data[name]
        st = stats[name]
        present = c["votes_za"] + c["votes_przeciw"] + c["votes_wstrzymal"] + c["votes_brak"]
        aktywnosc = (present / total_votes * 100) if total_votes else 0
        frekwencja = (len(st["sess"]) / total_sessions * 100) if total_sessions else 0
        total_decis = st["with"] + st["against"]
        zgodnosc = (st["with"] / total_decis * 100) if total_decis else 0.0
        councilors_list.append({
            "name": name, "club": c["club"], "district": None,
            "frekwencja": round(frekwencja, 1), "aktywnosc": round(aktywnosc, 1),
            "zgodnosc_z_klubem": round(zgodnosc, 1),
            "votes_za": c["votes_za"], "votes_przeciw": c["votes_przeciw"],
            "votes_wstrzymal": c["votes_wstrzymal"], "votes_brak": c["votes_brak"],
            "votes_nieobecny": c["votes_nieobecny"], "votes_total": total_votes,
            "rebellion_count": st["against"], "rebellions": [],
            "has_activity_data": False, "activity": None,
        })

    global NAME_AGG, _all_session_dates
    NAME_AGG = {name: dict(stats[name], sess=len(stats[name]["sess"])) for name in stats}
    _all_session_dates = [s["date"] for s in sessions_data]

    vectors = defaultdict(dict)
    for v in all_votes:
        for cat in ("za", "przeciw", "wstrzymal_sie"):
            for name in v["named_votes"].get(cat, []):
                vectors[name][v["id"]] = cat
    pairs = []
    names_sorted = sorted(vectors.keys())
    for a, b in combinations(names_sorted, 2):
        common = set(vectors[a].keys()) & set(vectors[b].keys())
        if len(common) < 10:
            continue
        same = sum(1 for vid in common if vectors[a][vid] == vectors[b][vid])
        score = round(same / len(common) * 100, 1)
        pairs.append({"a": a, "b": b, "club_a": _club_of(a), "club_b": _club_of(b),
                      "score": score, "common_votes": len(common)})
    pairs.sort(key=lambda x: x["score"], reverse=True)

    club_counts = Counter(_club_of(n) for n in all_names)
    kad = {
        "id": KADENCJA_ID, "label": KADENCJA_LABEL,
        "clubs": dict(club_counts),
        "sessions": sessions_data, "total_sessions": total_sessions,
        "total_votes": total_votes, "total_councilors": len(councilors_list),
        "councilors": councilors_list, "votes": all_votes,
        "similarity_top": pairs[:20], "similarity_bottom": pairs[-20:][::-1],
    }
    return {
        "generated": datetime.now().isoformat(),
        "default_kadencja": KADENCJA_ID,
        "kadencje": [kad],
    }


NAME_AGG = {}
_all_session_dates = []


def build_profiles(records):
    cv = defaultdict(lambda: {"za": 0, "przeciw": 0, "wstrzymal_sie": 0,
                              "nieobecny": 0, "brak": 0, "sess": set()})
    for rec in records:
        d = rec.get("session_date")
        if not d:
            continue
        for cat, names in rec["named"].items():
            for name in names:
                key = "za" if cat == "za" else "przeciw" if cat == "przeciw" \
                    else "wstrzymal_sie" if cat == "wstrzymal_sie" \
                    else "nieobecny" if cat == "nieobecni" else "brak"
                cv[name][key] += 1
                if key != "nieobecny":
                    cv[name]["sess"].add(d)
    profiles = []
    for name in sorted(cv.keys()):
        vd = cv[name]
        total = sum(vd[k] for k in ("za", "przeciw", "wstrzymal_sie", "nieobecny", "brak")) or 1
        agg = NAME_AGG.get(name, {})
        all_sess = len(vd["sess"])
        frekw = 100.0 * all_sess / len(_all_session_dates) if _all_session_dates else 0.0
        dec = agg.get("with", 0) + agg.get("against", 0)
        zgod = 100.0 * agg.get("with", 0) / dec if dec else 0.0
        profiles.append({
            "name": name, "slug": make_slug(name),
            "kadencje": {
                KADENCJA_ID: {
                    "club": _club_of(name), "has_voting_data": True,
                    "has_activity_data": False, "frekwencja": round(frekw, 1),
                    "aktywnosc": round(float(vd["za"] + vd["przeciw"] + vd["wstrzymal_sie"]) /
                                       total * 100, 1),
                    "zgodnosc_z_klubem": round(zgod, 1),
                    "votes_za": vd["za"], "votes_przeciw": vd["przeciw"],
                    "votes_wstrzymal": vd["wstrzymal_sie"], "votes_brak": vd["brak"],
                    "votes_nieobecny": vd["nieobecny"], "votes_total": total,
                    "rebellion_count": agg.get("against", 0), "rebellions": [],
                    "roles": [], "notes": "", "former": False, "mid_term": False,
                }
            }
        })
    return {"profiles": profiles}


def save_split(output, out_path, profiles):
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    stubs = []
    for kad in output.get("kadencje", []):
        kid = kad["id"]
        stubs.append({"id": kid, "label": kad.get("label", f"Kadencja {kid}")})
        with open(out_path.parent / f"kadencja-{kid}.json", "w", encoding="utf-8") as f:
            json.dump(kad, f, ensure_ascii=False, separators=(",", ":"))
    index = {"generated": output.get("generated", ""),
             "default_kadencja": output.get("default_kadencja", ""),
             "kadencje": stubs}
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(index, f, ensure_ascii=False, separators=(",", ":"))
    with open(out_path.parent / "profiles.json", "w", encoding="utf-8") as f:
        json.dump(profiles, f, ensure_ascii=False, separators=(",", ":"))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--output", required=True, help="docs/data.json")
    ap.add_argument("--profiles", required=True, help="docs/profiles.json")
    ap.add_argument("--cache-dir", default=None)
    args = ap.parse_args()
    cache_dir = Path(args.cache_dir) if args.cache_dir else None

    print(f"=== Scraper Rada Miasta Piotrków Trybunalski ({BIP}) ===")
    sessions = collect_sessions(cache_dir)
    print(f"  Sesji z opublikowanymi wynikami DOCX: {len(sessions)}")

    records = collect_all(sessions, cache_dir)
    print(f"  Razem glosowań imiennych: {len(records)}")

    for r in records:
        _canonical_rec(r)

    if not records:
        print("  BRAK DANYCH — nic do zapisania.")
        sys.exit(1)

    output = build_output(records)
    profiles = build_profiles(records)
    save_split(output, args.output, profiles)
    total = output["kadencje"][0]
    print(f"  Zapisano: {args.output}, kadencja-{KADENCJA_ID}.json, profiles.json")
    print(f"  Sesji: {total['total_sessions']}, glosowan: {total['total_votes']}, "
          f"radnych: {total['total_councilors']}")


if __name__ == "__main__":
    main()
