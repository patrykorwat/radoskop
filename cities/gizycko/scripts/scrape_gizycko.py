#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Radoskop Giżycko — imienne głosowania Rady Miasta Giżycka (IX kadencja).

Źródło: BIP bip.gizycko.pl (platforma idcom), Rada Miasta → kategoria
"Protokoły sesji i wyniki głosowań Rady Miasta" → artykuł zbiorczy
"WYNIKI GŁOSOWAŃ RADNYCH PODCZAS SESJI VIII i IX KADENCJI..." (uid 455850),
który zawiera per-sesję załączniki w 3 formatach (mieszanka zależna od roku):

  A) RAR5/ZIP archiwum (2024-05 … 2025-09): PDF per głosowanie w formacie
     wydruku eSesja — dwukolumnowa tabela "LP | NAZWISKO I IMIĘ | GŁOS"
     (za/przeciw/wstrzymał się/nieobecny/obecny), agregaty "GŁOSY ZA/PRZECIW/
     WSTRZYMUJĄCE SIĘ", walidowalne per głos. RAR czytany przez lib_archive
     (ctypes + systemowy libarchive, bez CLI unrar).
  B) DOCX raport (2025-11 … 2026-03): "Raport z głosowań" — tekst z blokami
     "Głosowano w sprawie: … / Wyniki głosowania: ZA: n, … / Wyniki imienne:
     ZA (n) imienna lista po przecinku …". parser esesja-TEXT (wzorowany na
     referencji esesja-imienne-text-format).
  C) PDF raport (2025-10 … 2026-07): ten sam layout co DOCX jako PDF —
     części mają warstwę tekstową, część to skany (OCR tesseract -l pol).

Data sesji: z nazwy załącznika (dd.mm.rrrr); dla archiwów dodatkowa walidacja
datą "DATA GŁOSOWANIA" pierwszego PDF. Nazwa tematu: z bloku głosowania.

Użycie:
    python scrape_gizycko.py --city-dir <cities/gizycko> [--cache-dir dir]
Zapisuje: docs/data.json, docs/kadencja-2024-2029.json, docs/profiles.json
"""
import argparse
import hashlib
import io
import json
import os
import re
import subprocess
import sys
import tempfile
import time
import unicodedata
import zipfile
import zlib
from collections import Counter, defaultdict
from datetime import datetime, date
from pathlib import Path

import pdfplumber
import requests
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

sys.path.insert(0, str(Path(__file__).resolve().parents[3] / "scripts"))
from lib_archive import read_archive_bytes  # noqa: E402

BIP = "https://bip.gizycko.pl"
VOTES_ARTICLE = ("/organy/106/dokumenty/1224/wiadomosc/455850/"
                 "wyniki_glosowan_radnych_podczas_sesji_viii_i_ix_kadencji_rady_mi")
KAD_START = "2024-05-07"
KADENCJA_ID = "2024-2029"
KADENCJA_LABEL = "IX kadencja (2024\u20132029)"

_MONTHS = {
    "stycznia":1,"lutego":2,"marca":3,"kwietnia":4,"maja":5,"czerwca":6,"lipca":7,
    "sierpnia":8,"września":9,"pazdziernika":10,"października":10,"listopada":11,"grudnia":12,
    "styczeń":1,"luty":2,"marzec":3,"kwiecień":4,"maj":5,"czerwiec":6,"lipiec":7,
    "sierpień":8,"wrzesień":9,"pazdziernik":10,"październik":10,"listopad":11,"grudzień":12,
}

def _nk(s):
    s = s.lower().replace("ł", "l")
    s = unicodedata.normalize("NFKD", s)
    return re.sub(r"[^a-z0-9]", "", "".join(c for c in s if not unicodedata.combining(c)))

def _norm_vote(w):
    k = _nk(w)
    if k in ("za", "z", "ża", "ze"):
        return "za"
    if k.startswith("przeciw"):
        return "przeciw"
    if k.startswith("wstrzym"):
        return "wstrzymal_sie"
    if k.startswith("nieobecn"):
        return "nieobecni"
    if k.startswith("obecn"):
        return "obecny"
    if k in ("brakglosu", "brak"):
        return "brak_glosu"
    return None

REQ_DELAY = 0.6
_LAST = 0.0

def _rate():
    global _LAST
    d = time.time() - _LAST
    if d < REQ_DELAY:
        time.sleep(REQ_DELAY - d)
    _LAST = time.time()

def _get(url, cache_dir, tries=4):
    key = hashlib.md5(url.encode()).hexdigest()
    cf = None
    if cache_dir:
        cdir = Path(cache_dir)
        cdir.mkdir(parents=True, exist_ok=True)
        cf = cdir / (key + ".dat")
        if cf.is_file():
            return cf.read_bytes()
    last = None
    for a in range(tries):
        try:
            _rate()
            r = requests.get(url, headers={"User-Agent": "Mozilla/5.0 (Radoskop)"}, timeout=90, verify=False)
            r.raise_for_status()
            data = r.content
            if cf is not None:
                cf.write_bytes(data)
            return data
        except Exception as e:  # noqa: BLE001
            last = e
            time.sleep(1.5 * (a + 1))
    raise RuntimeError(f"GET fail {url}: {last}")

# ---------------- discovery ----------------
def discover_attachments():
    """Zwraca listę {date, name, url} załączników IX kadencji."""
    html_doc = _get(BIP + VOTES_ARTICLE, None).decode("utf-8", "ignore")
    urls = sorted(set(re.findall(r'href="(https://bip-v1-files\.idcom-jst\.pl/[^"]+)"', html_doc)))
    out = []
    for u in urls:
        n = u.rsplit("/", 1)[-1]
        if not re.search(r"osow|sesja", n, re.I):  # 'gosowanie' — literówki w nazwach BIP
            continue
        m = re.search(r"(\d{2})[._-]?(\d{2})[._-]?(\d{4})", n)
        if not m:
            continue
        try:
            d = date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
        except ValueError:
            continue
        iso = d.isoformat()
        if iso < KAD_START:
            continue
        if not n.lower().endswith((".rar", ".zip", ".pdf", ".docx")):
            continue
        out.append({"date": iso, "name": n, "url": u})
    # dedup po dacie (ten sam dzień potrafi mieć duplikat wariantu)
    seen = {}
    for a in out:
        seen.setdefault(a["date"], a)
    res = sorted(seen.values(), key=lambda x: x["date"])
    return res

# ---------------- per-vote PDF table parser (wzorowany na goleniow) ----------------
def _table_region(words):
    up = [w for w in words if _nk(w["text"]) == "uprawnieni"]
    if not up:
        return words, 0
    thr = max(w["top"] for w in up)
    return [w for w in words if w["top"] > thr + 4], thr

def _col_boundary(words):
    right_lps = [w["x0"] for w in words if re.match(r"^\d{1,2}\.?$", w["text"]) and w["x0"] > 280]
    if right_lps:
        return min(right_lps) - 3
    lps = [w["x0"] for w in words if re.match(r"^\d{1,2}\.?$", w["text"])]
    return (min(lps) + max(lps)) / 2.0 if lps else 297.0

def _parse_column(col):
    col.sort(key=lambda t: (t[0], t[1]))
    rows = []
    cr = None
    for top, x0, t in col:
        if cr is None or top - cr[0] > 6:
            cr = [top, []]
            rows.append(cr)
        cr[1].append((x0, t))
    out = []
    cur = None
    def emit():
        nonlocal cur
        if cur is not None and cur.get("vote"):
            out.append((cur["name"].strip(), cur["vote"]))
        cur = None
    for top, toks in rows:
        toks.sort(key=lambda z: z[0])
        if re.match(r"^\d+\.?$", toks[0][1]):
            emit()
            cur = {"name": "", "vote": None}
            toks = toks[1:]
        elif cur is None:
            cur = {"name": "", "vote": None}
        for _x, t in toks:
            nv = _norm_vote(t)
            if nv:
                cur["vote"] = nv
            elif _nk(t) in ("sie", "sier"):
                pass
            elif re.match(r"(?i)^(wydrukowano:?|\d{1,2}[:.]\d{2}([:.]\d{2})?$|\d{1,2}\.\d{1,2}\.\d{4})", t):
                pass
            else:
                cur["name"] = (cur["name"] + " " + t).strip()
    emit()
    return out

def _table_cells(words):
    b = _col_boundary(words)
    L = [(w["top"], w["x0"], w["text"]) for w in words if w["x0"] < b]
    R = [(w["top"], w["x0"], w["text"]) for w in words if w["x0"] >= b]
    cells = []
    for col in (L, R):
        cells += _parse_column(col)
    return cells

def _extract_aggs(text):
    agg = {}
    t = text.upper()
    for key, pat in [("uprawnionych", r"UPRAWNIONYCH\s+(\d+)"),
                     ("obecnych", r"OBECNYCH\s+(\d+)"),
                     ("nieobecnych", r"NIEOBECNYCH\s+(\d+)"),
                     ("za", r"ZA\s+(\d+)"),
                     ("przeciw", r"PRZECIW\s+(\d+)"),
                     ("wstrzym", r"WSTRZYMUJ[AĆĄ]CE SI[ĘE]\s+(\d+)"),
                     ("nieoddane", r"NIEODDANE\s+(\d+)")]:
        m = re.search(pat, t)
        if m:
            agg[key] = int(m.group(1))
    return agg

def _extract_topic_table(text):
    m = re.search(r"(?i)TYP G\u0141OSOWANIA", text)
    pre = text[:m.start()] if m else text
    lines = [l.strip() for l in pre.split("\n")]
    out = []
    for l in lines:
        if not l:
            continue
        if re.search(r"(?i)G\u0141OSOWANIE|Sesja|kadencj", l):
            continue
        if re.match(r"^(\d+\.?|Nr\s*\d+\.?)$", l):
            continue
        out.append(l)
    topic = " ".join(out)
    topic = re.sub(r"\s+", " ", topic).strip(" .:,;-")
    return topic or "(glosowanie)"

def parse_vote_pdf_bytes(data):
    """PDF per-głosowanie (tabela dwukolumnowa) → lista z jednym votem lub []"""
    recs = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        cur = None
        for page in pdf.pages:
            words = page.extract_words()
            text = page.extract_text() or ""
            has_agg = bool(re.search(r"(?i)LICZBA UPRAWNIONYCH", text))
            tw, _ = _table_region(words)
            cells = _table_cells(tw)
            if has_agg:
                agg = _extract_aggs(text)
                topic = _extract_topic_table(text)
                cur = {"topic": topic, "agg": agg, "cells": cells}
                recs.append(cur)
            elif cur is not None:
                cur["cells"] += cells
    votes = []
    for v in recs:
        counter = Counter(vote for _n, vote in v["cells"])
        agg = v["agg"]
        ok = (counter.get("za", 0) == agg.get("za", -1) and
              counter.get("przeciw", 0) == agg.get("przeciw", -1) and
              counter.get("wstrzymal_sie", 0) == agg.get("wstrzym", -1))
        if "nieobecnych" in agg:
            ok = ok and counter.get("nieobecni", 0) == agg["nieobecnych"]
        named = {"za": [], "przeciw": [], "wstrzymal_sie": [], "nieobecni": []}
        for name, vote in v["cells"]:
            if vote in named:
                named[vote].append(display_name(name))
        if ok:
            votes.append({"topic": v["topic"], "named": named})
    return votes

def display_name(surname_first):
    """'Andruszkiewicz Paweł Michał' -> 'Paweł Michał Andruszkiewicz' (nazwisko na koniec)."""
    toks = surname_first.split()
    if len(toks) >= 2:
        name = " ".join(toks[1:] + [toks[0]])
    else:
        name = surname_first
    # OCR/line-wrap artefakty: 'Koroś- Mieronowicz' -> 'Koroś-Mieronowicz'
    name = re.sub(r"([a-ząćęłńóśźż])-\s+([A-ZĄĆĘŁŃÓŚŹŻ])", r"\1-\2", name)
    name = re.sub(r"\s+", " ", name).strip()
    # OCR literówki known-in-roster wariantów bez ogonków
    fixed = {"Pawel Łazarski": "Paweł Łazarski"}
    return fixed.get(name, name)

# ---------------- raport text/DOCX parser (esesja TEXT layout) ----------------
_REPORT_HDR = re.compile(r"(?i)W.{2,3}ki g\u0142osowania.*?\n?\s*ZA:\s*(\d+),\s*PRZECIW:\s*(\d+),\s*WSTRZYMUJ[ĘE] SI[ĘE]:\s*(\d+),\s*BRAK G\u0141OSU:\s*(\d+),\s*NIEOBECNI:\s*(\d+)")
_ROLL_HDR = re.compile(r"(?i)LICZBA UPRAWNIONYCH")

_FOOTER_MARKERS = ["czas głosowania", "uczestnictwo w głosowaniach", "wygenerowano",
                   "głosowanie z dnia", "wyniki głosowania (radni)", "wyniki głosowania:"]

_OCR_NAME_FIX = {"Pawel": "Paweł", "Piotr Maciej Tomanek": "Piotr Maciej Tomanek"}

def _fix_name(t):
    """Normalizacja nazwiska z raportu: zrost po łamaniu wiersza po myślniku + OCR bez ogonków."""
    t = re.sub(r"([a-ząćęłńóśźż])-\s+([A-ZĄĆĘŁŃÓŚŹŻ])", r"\1-\2", t)
    t = re.sub(r"\s+", " ", t).strip()
    for bad, good in _OCR_NAME_FIX.items():
        if t.startswith(bad) and bad != good:
            t = good + t[len(bad):]
    return t

def _clean_lists(chunk):
    chunk = re.sub(r"\s+", " ", chunk).strip()
    for fm in _FOOTER_MARKERS:
        i = chunk.lower().find(fm)
        if i != -1:
            chunk = chunk[:i]
    names = []
    for tok in chunk.split(","):
        tok = tok.strip(" .;")
        if not tok:
            continue
        if not re.search(r"[A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż]", tok):
            continue
        if re.search(r"(?i)^\d+$", tok):
            continue
        names.append(tok)
    return names

def parse_report_text(text, ocr_fallback_pages=None):
    """Blokowe raporty (DOCX / PDF tekstowy): bloki 'Głosowano (w sprawie|wniosek...)'.

    Zwraca listę {topic, named} z walidacją sum vs header counts.
    """
    votes = []
    # split into blocks by the vote header marker
    parts = re.split(r"(?=\n\s*\d+\.\s*G\u0142osowano\b)", "\n" + text)
    for part in parts:
        m = re.match(r"\s*\d+\.\s*(G\u0142osowano [^\n]*)", part)
        if not m:
            continue
        hm = _REPORT_HDR.search(part)
        if not hm:
            continue
        za, pr, ws, brak, nieob = (int(hm.group(i)) for i in range(1, 6))
        # topic: po 'Głosowano ... w sprawie:' do 'czas głosowania' / 'Wyniki'
        seg = part[:hm.start()]
        tm = re.search(r"(?i)(?:G\u0142osowano (?:w sprawie|wniosek w sprawie|projektu w sprawie|nad|w sprawie wniosku)?[:\s]*)", seg)
        topic = seg[tm.end():] if tm else seg
        topic = re.split(r"(?i)czas g\u0142osowania", topic)[0]
        topic = re.sub(r"\s+", " ", topic).strip(" .:,;-")
        im = re.search(r"(?i)Wyniki imienne\s*:?", part)
        if not im:
            continue
        chunk = part[im.end():]
        labels = list(re.finditer(r"(?i)^(ZA|PRZECIW|WSTRZYMUJ[ĘE] SI[ĘE]|BRAK G\u0141OSU|NIEOBECNI)\s*\(\d+\)", chunk, re.M))
        if not labels:
            continue
        cats = {}
        for i, lm in enumerate(labels):
            end = labels[i + 1].start() if i + 1 < len(labels) else len(chunk)
            body = chunk[lm.end():end]
            key = _norm_vote(lm.group(1)) or "za"
            key = {"wstrzymal_sie": "wstrzymal_sie", "brak_glosu": "brak_glosu", "nieobecni": "nieobecni"}.get(key, key)
            cats[key] = [_fix_name(t) for t in _clean_lists(body)]
        got = {k: len(v) for k, v in cats.items()}
        want = {"za": za, "przeciw": pr, "wstrzymal_sie": ws, "brak_glosu": brak, "nieobecni": nieob}
        if any(got.get(k, 0) != n for k, n in want.items()):
            continue
        named = {k: cats.get(k, []) for k in ("za", "przeciw", "wstrzymal_sie", "nieobecni")}
        votes.append({"topic": topic or "(glosowanie)", "named": named})
    return votes

def parse_docx_bytes(data):
    import docx
    d = docx.Document(io.BytesIO(data))
    lines = [p.text for p in d.paragraphs]
    for t in d.tables:  # raport giżycki nie ma tabel, ale defensively
        for row in t.rows:
            for c in row.cells:
                lines.append(c.text)
    return parse_report_text("\n".join(lines))

def pdf_text_with_ocr(data):
    """Warstwa tekstowa + OCR dla stron bez tekstu (skany). OCR po jednej stronie,
    sekwencyjnie (nie równolegle — pitfall znany)."""
    import pymupdf
    doc = pymupdf.open(stream=data, filetype="pdf")
    out = []
    for i, page in enumerate(doc):
        t = page.get_text()
        if len(t.strip()) < 40:
            pix = page.get_pixmap(dpi=150)
            png = pix.tobytes("png")
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
                tf.write(png)
                pth = tf.name
            try:
                r = subprocess.run(["tesseract", pth, "-", "-l", "pol"], capture_output=True, text=True, timeout=180)
                t = r.stdout
            except Exception:
                t = ""
            finally:
                try:
                    os.unlink(pth)
                except OSError:
                    pass
        out.append(t)
    doc.close()
    return "\n".join(out)

def parse_report_pdf(data):
    text = pdf_text_with_ocr(data)
    votes = parse_report_text(text)
    if votes:
        return votes
    return parse_vote_pdf_bytes(data)  # fallback: single-vote table PDF

# ---------------- attach dispatch ----------------
def votes_from_attachment(att, cache):
    """Per-załącznik → lista {date, topic, named}."""
    blob = _get(att["url"], cache)
    name = att["name"].lower()
    out = []
    if name.endswith(".docx"):
        out = parse_docx_bytes(blob)
    elif name.endswith(".pdf"):
        # duży raport wielogłosowy vs pojedynczy PDF głosowania
        if blob[:4] == b"%PDF":
            head = ""
            try:
                with pdfplumber.open(io.BytesIO(blob)) as pdf:
                    head = (pdf.pages[0].extract_text() or "")[:200]
            except Exception:
                head = ""
            if "Raport" in head or "Wyniki imienne" in head:
                out = parse_report_pdf(blob)
            else:
                # najpierw raport tekstowy (OCR), potem tabela dwukolumnowa
                out = parse_report_pdf(blob)
    elif name.endswith((".rar", ".zip")):
        try:
            ents = read_archive_bytes(blob)
        except RuntimeError:
            ents = {}
        if not ents:
            try:
                with zipfile.ZipFile(io.BytesIO(blob)) as z:
                    ents = {n: z.read(n) for n in z.namelist() if not n.endswith("/")}
            except Exception:
                ents = {}
        for en, edata in sorted(ents.items()):
            if not en.lower().endswith(".pdf"):
                continue
            try:
                vs = parse_vote_pdf_bytes(edata)
            except Exception:
                vs = []
            out.extend(vs)
    for v in out:
        v["date"] = att["date"]
    return out

def display_named(v):
    named = {k: [display_name(n) for n in v] for k, v in v["named"].items()}
    return {"topic": v["topic"], "named": named}

# ---------------- output (wzorzec goleniow) ----------------
def make_slug(name):
    repl = {'ą':'a','ć':'c','ę':'e','ł':'l','ń':'n','ó':'o','ś':'s','ź':'z','ż':'z',
            'Ą':'A','Ć':'C','Ę':'E','Ł':'L','Ń':'N','Ó':'O','Ś':'S','Ź':'Z','Ż':'Z'}
    slug = name.lower()
    for pl, a in repl.items():
        slug = slug.replace(pl, a)
    return re.sub(r"[^a-z0-9]+", "", slug)

def build(records, club_assign=None):
    club_assign = club_assign or {}
    all_votes = []
    vid = 0
    sessions_by_date = {}
    for rec in records:
        d = rec["date"]
        if not d or d < KAD_START:
            continue
        if d not in sessions_by_date:
            sessions_by_date[d] = {"date": d, "number": "", "vote_count": 0, "attendees": set(), "speakers": []}
        sessions_by_date[d]["vote_count"] += 1
        named = {k: list(v) for k, v in rec["named"].items()}
        for cat in ("za", "przeciw", "wstrzymal_sie", "brak_glosu"):
            sessions_by_date[d]["attendees"].update(rec["named"].get(cat, []))
        vid += 1
        all_votes.append({"id": str(vid), "session_date": d, "session_number": "",
                          "topic": rec.get("topic", ""), "named_votes": named,
                          "counts": {k: len(named.get(k, [])) for k in ("za", "przeciw", "wstrzymal_sie")}})
    sessions_data = []
    for d in sorted(sessions_by_date.keys()):
        s = sessions_by_date[d]
        sessions_data.append({"date": d, "number": s["number"], "vote_count": s["vote_count"],
                              "attendee_count": len(s["attendees"]), "attendees": sorted(s["attendees"]),
                              "speakers": []})
    all_names = set()
    for v in all_votes:
        for names in v["named_votes"].values():
            all_names.update(names)
    councilors_data = {}
    for name in sorted(all_names):
        councilors_data[name] = {"name": name, "club": club_assign.get(name, "NZ"), "district": None,
                                 "votes_za": 0, "votes_przeciw": 0, "votes_wstrzymal": 0,
                                 "votes_brak": 0, "votes_nieobecny": 0, "rebellions": []}
    for v in all_votes:
        for cat, names in v["named_votes"].items():
            if cat == "nieobecni":
                for nm in names:
                    if nm in councilors_data:
                        councilors_data[nm]["votes_nieobecny"] += 1
                continue
            if cat == "brak_glosu":
                for nm in names:
                    if nm in councilors_data:
                        councilors_data[nm]["votes_brak"] += 1
                continue
            for nm in names:
                if nm not in councilors_data:
                    continue
                if cat == "za":
                    councilors_data[nm]["votes_za"] += 1
                elif cat == "przeciw":
                    councilors_data[nm]["votes_przeciw"] += 1
                else:
                    councilors_data[nm]["votes_wstrzymal"] += 1
    total_votes = len(all_votes)
    total_sessions = len(sessions_data)
    councillor_sess = defaultdict(set)
    for v in all_votes:
        for cat, names in v["named_votes"].items():
            if cat == "nieobecni":
                continue
            for nm in names:
                councillor_sess[nm].add(v["session_date"])
    councilors_list = []
    for c in sorted(councilors_data.values(), key=lambda x: x["name"]):
        present = c["votes_za"] + c["votes_przeciw"] + c["votes_wstrzymal"]
        aktywn = (present / total_votes * 100) if total_votes else 0
        frekw = (len(councillor_sess.get(c["name"], set())) / total_sessions * 100) if total_sessions else 0
        councilors_list.append({"name": c["name"], "club": c["club"], "district": None,
                                "frekwencja": round(frekw, 1), "aktywnosc": round(aktywn, 1),
                                "zgodnosc_z_klubem": 0.0,
                                "votes_za": c["votes_za"], "votes_przeciw": c["votes_przeciw"],
                                "votes_wstrzymal": c["votes_wstrzymal"], "votes_brak": c["votes_brak"],
                                "votes_nieobecny": c["votes_nieobecny"], "votes_total": total_votes,
                                "rebellion_count": 0, "rebellions": [],
                                "has_activity_data": False, "activity": None})
    vectors = defaultdict(dict)
    for v in all_votes:
        for cat in ("za", "przeciw", "wstrzymal_sie"):
            for nm in v["named_votes"].get(cat, []):
                vectors[nm][v["id"]] = cat
    from itertools import combinations
    pairs = []
    names_sorted = sorted(vectors.keys())
    for a, b in combinations(names_sorted, 2):
        common = set(vectors[a].keys()) & set(vectors[b].keys())
        if len(common) < 10:
            continue
        same = sum(1 for vid_ in common if vectors[a][vid_] == vectors[b][vid_])
        pairs.append({"a": a, "b": b, "club_a": "", "club_b": "",
                      "score": round(same / len(common) * 100, 1), "common_votes": len(common)})
    pairs.sort(key=lambda x: x["score"], reverse=True)
    kad = {"id": KADENCJA_ID, "label": KADENCJA_LABEL,
           "clubs": dict(Counter(club_assign.get(c["name"], "NZ") for c in councilors_list)),
           "sessions": sessions_data, "total_sessions": total_sessions,
           "total_votes": total_votes, "total_councilors": len(councilors_list),
           "councilors": councilors_list, "votes": all_votes,
           "similarity_top": pairs[:20], "similarity_bottom": pairs[-20:][::-1]}
    return {"generated": datetime.now().isoformat(), "default_kadencja": KADENCJA_ID, "kadencje": [kad]}, total_votes, total_sessions

def build_profiles(records, total_votes, club_assign=None):
    club_assign = club_assign or {}
    cv = defaultdict(lambda: {"za": 0, "przeciw": 0, "wstrzymal_sie": 0, "nieobecni": 0, "votes": []})
    for rec in records:
        d = rec["date"]
        if not d or d < KAD_START:
            continue
        for cat, names in rec["named"].items():
            for nm in names:
                if cat in ("za", "przeciw", "wstrzymal_sie", "nieobecni"):
                    cv[nm][cat] += 1
                cv[nm]["votes"].append({"session": d, "vote": cat})
    sess_set = {r["date"] for r in records if r["date"] >= KAD_START}
    n_sessions = len(sess_set) or 1
    profiles = []
    for nm in sorted(cv.keys()):
        vd = cv[nm]
        total = sum(vd[k] for k in ("za", "przeciw", "wstrzymal_sie")) or 1
        sess = len({v["session"] for v in vd["votes"] if v["vote"] != "nieobecni"})
        aktywn = (vd["za"] + vd["przeciw"] + vd["wstrzymal_sie"]) / n_sessions * 100
        profiles.append({"name": nm, "slug": make_slug(nm),
                         "kadencje": {KADENCJA_ID: {"club": club_assign.get(nm, "NZ"),
                             "has_voting_data": True, "has_activity_data": False,
                             "frekwencja": round(sess / n_sessions * 100, 1),
                             "aktywnosc": round(aktywn, 1), "zgodnosc_z_klubem": 0.0,
                             "votes_za": vd["za"], "votes_przeciw": vd["przeciw"],
                             "votes_wstrzymal": vd["wstrzymal_sie"], "votes_brak": 0,
                             "votes_nieobecny": vd["nieobecni"], "votes_total": total,
                             "rebellion_count": 0, "rebellions": [], "roles": [],
                             "notes": "", "former": False, "mid_term": False}}})
    return {"profiles": profiles, "total": len(profiles)}

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--city-dir", required=True)
    ap.add_argument("--cache-dir", default=None)
    args = ap.parse_args()
    city_dir = Path(args.city_dir)
    cache = Path(args.cache_dir) if args.cache_dir else None
    cfg = {}
    if (city_dir / "config.json").is_file():
        cfg = json.loads((city_dir / "config.json").read_text(encoding="utf-8"))
    club_assign = cfg.get("club_assignments", {}) or {}

    atts = discover_attachments()
    print(f"[gizycko] załaczników IX kad: {len(atts)}")
    records = []
    for att in atts:
        try:
            vs = votes_from_attachment(att, cache)
            records += vs
            print(f"  [ok] {att['date']} {att['name'][:52]:52s} votes={len(vs)}")
        except Exception as e:
            print(f"  [ERR {att['date']} {att['name'][:40]}] {type(e).__name__}: {e}")
    # głosy z tego samego dnia z kilku załączników (rzadkość) zostają; sort
    records.sort(key=lambda r: r["date"])
    output, total_votes, total_sessions = build(records, club_assign)
    profiles = build_profiles(records, total_votes, club_assign)
    docs = city_dir / "docs"
    docs.mkdir(exist_ok=True)
    (docs / "kadencja-2024-2029.json").write_text(
        json.dumps(output["kadencje"][0], ensure_ascii=False, indent=1), encoding="utf-8")
    data = {"generated": output["generated"], "default_kadencja": KADENCJA_ID,
            "kadencje": [{"id": KADENCJA_ID, "label": KADENCJA_LABEL}]}
    (docs / "data.json").write_text(json.dumps(data, ensure_ascii=False, indent=1), encoding="utf-8")
    (docs / "profiles.json").write_text(json.dumps(profiles, ensure_ascii=False, indent=1), encoding="utf-8")
    print(f"[gizycko] DONE votes={total_votes} sessions={total_sessions} councilors={profiles['total']}")

if __name__ == "__main__":
    main()
